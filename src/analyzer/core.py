import os
import io
import re
import PyPDF2
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from dotenv import load_dotenv

load_dotenv()

# Configure Groq
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

class ResumeAnalyzer:
    def __init__(self, model_name="llama-3.3-70b-versatile"):
        self.model = model_name

    def extract_text_from_pdf(self, pdf_path):
        """Extracts text from a local PDF file."""
        try:
            with open(pdf_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    async def analyze(self, resume_text, jd_text):
        """Uses Groq to analyze the resume against the JD with a fun, emoji-rich, interactive scorecard."""
        prompt = f"""
        Act as a Friendly Principal Talent Scout. 
        Analyze the Resume against the Job Description (JD) using lots of emojis.
        
        JD: {jd_text}
        Resume: {resume_text}

        Response format (Use these EXACT headers):
        📊 **MATCH SCORECARD**
        Rating: [X/10] 🚀
        
        ⚠️ **DANGER ZONE: CRITICAL GAPS**
        [List what is in the JD but missing in the Resume - be specific!]

        💡 **OPTIMIZATION POWER-UPS**
        [3 specific hacks to make the resume better]

        🗺️ **CAREER PATHWAY RADAR**
        [Short, encouraging professional advice]
        """
        try:
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=self.model,
            )
            return chat_completion.choices[0].message.content
        except Exception as e:
            return f"❌ Error during Groq analysis: {str(e)}"

    async def generate_updated_resume(self, resume_text, jd_text):
        """Rewrites the resume to match the JD using Groq, removing markdown artifacts."""
        prompt = f"""
        Expert Executive Resume Writer: Rewrite this resume to perfectly match the Job Description.
        Use a high-impact professional tone. 

        RULES:
        1. DO NOT use markdown symbols like **, __, or #.
        2. Ensure the full name is the first line.
        3. Organize into: Summary, Skills, Work Experience, Education, Projects.
        4. Focus on quantified achievements (e.g., "Increased X by 20%").
        
        JD: {jd_text}
        Original CV: {resume_text}

        Output only the updated resume text.
        """
        try:
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=self.model,
            )
            clean_text = chat_completion.choices[0].message.content
            # Remove any residual markdown characters just in case
            clean_text = re.sub(r'[*#_]', '', clean_text)
            return clean_text
        except Exception as e:
            return f"Error: {str(e)}"

    def generate_docx(self, text):
        """Converts text into a PREMIUM STYLE Word document."""
        doc = Document()
        
        # Clean the text
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            return None

        # 1. Header (Name)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(lines[0].upper())
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_run.font.name = 'Calibri'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. Body
        current_section = None
        sections = ["SUMMARY", "SKILLS", "WORK EXPERIENCE", "EXPERIENCE", "EDUCATION", "PROJECTS", "TRAINING"]

        for line in lines[1:]:
            upper_line = line.upper()
            is_section = any(s in upper_line for s in sections) and len(line) < 25
            
            if is_section:
                doc.add_paragraph() # Spacer
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(upper_line)
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue accent
                p.paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                if line.startswith('-') or line.startswith('•'):
                    p.style = 'List Bullet'
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def generate_pdf(self, text):
        """Converts text into a clean PDF document with basic styling."""
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if not lines:
            return None

        # Header
        c.setFont("Helvetica-Bold", 18)
        c.drawCentredString(width/2, height - 50, lines[0].upper())
        
        c.setFont("Helvetica", 10)
        y = height - 80
        
        sections = ["SUMMARY", "SKILLS", "WORK EXPERIENCE", "EXPERIENCE", "EDUCATION", "PROJECTS", "TRAINING"]

        for line in lines[1:]:
            upper_line = line.upper()
            is_section = any(s in upper_line for s in sections) and len(line) < 25

            if is_section:
                y -= 25
                c.setFont("Helvetica-Bold", 12)
                c.drawString(40, y, upper_line)
                c.line(40, y-2, width-40, y-2) # Horizontal line
                y -= 15
                c.setFont("Helvetica", 10)
            else:
                # Basic wrapping
                text_obj = c.beginText(40, y)
                text_obj.setFont("Helvetica", 10)
                words = line.split()
                line_to_draw = ""
                for word in words:
                    if c.stringWidth(line_to_draw + " " + word, "Helvetica", 10) < (width - 80):
                        line_to_draw += " " + word if line_to_draw else word
                    else:
                        text_obj.textLine(line_to_draw)
                        line_to_draw = word
                text_obj.textLine(line_to_draw)
                c.drawText(text_obj)
                y = text_obj.getY() - 12
            
            if y < 50:
                c.showPage()
                y = height - 50
        
        c.save()
        buffer.seek(0)
        return buffer
